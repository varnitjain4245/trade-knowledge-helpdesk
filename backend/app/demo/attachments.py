"""Attachments for the assistant.

Lets a person put a document in front of the desk and ask about it — "what does this
circular say about my product" is the question the whole helpdesk exists to answer, and
until now it could only be asked about records already indexed.

Extracted text becomes retrieval context for that conversation only. It is deliberately
**not** added to the knowledge base: a document someone uploaded is not an approved
record, and letting it answer other people's questions would break the approval chain
that makes citations mean anything.
"""

from __future__ import annotations

import io
from dataclasses import dataclass, field
from datetime import datetime, timezone

TEXT_TYPES = {"text/plain", "text/markdown", "text/csv", "application/json"}
PDF_TYPES = {"application/pdf"}
DOCX_TYPES = {
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}
IMAGE_TYPES = {"image/png", "image/jpeg", "image/webp", "image/gif"}

MAX_BYTES = 8 * 1024 * 1024
MAX_CHARS = 24_000


@dataclass
class Attachment:
    name: str
    content_type: str
    byte_size: int
    text: str
    pages: int | None = None
    note: str = ""
    at: datetime = field(default_factory=lambda: datetime.now(timezone.utc))

    @property
    def usable(self) -> bool:
        return bool(self.text.strip())


class ExtractionFailed(Exception):
    """Raised with a reason a person can act on, never a bare failure."""


def extract(name: str, content_type: str, payload: bytes) -> Attachment:
    if len(payload) > MAX_BYTES:
        raise ExtractionFailed(
            f"That file is {len(payload) // (1024 * 1024)} MB. The limit is "
            f"{MAX_BYTES // (1024 * 1024)} MB — send the relevant pages instead."
        )
    if not payload:
        raise ExtractionFailed("That file was empty.")

    kind = (content_type or "").split(";")[0].strip().lower()
    lower = name.lower()

    if kind in PDF_TYPES or lower.endswith(".pdf"):
        return _pdf(name, kind or "application/pdf", payload)
    if kind in DOCX_TYPES or lower.endswith(".docx"):
        return _docx(name, kind, payload)
    if kind in TEXT_TYPES or lower.endswith((".txt", ".md", ".csv", ".json")):
        return _plain(name, kind or "text/plain", payload)
    if kind in IMAGE_TYPES or lower.endswith((".png", ".jpg", ".jpeg", ".webp")):
        # Said plainly rather than returning an empty result that looks like a failure to
        # understand the document. No vision or OCR service is configured, and guessing
        # at what an image says would be exactly the invention this product forbids.
        raise ExtractionFailed(
            "Images cannot be read yet — no text-recognition service is configured. "
            "Send the document as a PDF or paste the passage you are asking about."
        )
    raise ExtractionFailed(
        f"{name} is not a file type the desk can read. PDF, Word, plain text and CSV work."
    )


def _plain(name: str, kind: str, payload: bytes) -> Attachment:
    for encoding in ("utf-8", "utf-16", "latin-1"):
        try:
            text = payload.decode(encoding)
            break
        except UnicodeDecodeError:
            continue
    else:
        raise ExtractionFailed("That file's text could not be decoded.")
    return Attachment(name, kind, len(payload), text[:MAX_CHARS])


def _pdf(name: str, kind: str, payload: bytes) -> Attachment:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise ExtractionFailed("PDF reading is unavailable on this server.") from exc

    try:
        reader = PdfReader(io.BytesIO(payload))
    except Exception as exc:
        raise ExtractionFailed("That PDF could not be opened — it may be damaged.") from exc

    if reader.is_encrypted:
        raise ExtractionFailed("That PDF is password-protected.")

    parts = []
    for page in reader.pages:
        try:
            parts.append(page.extract_text() or "")
        except Exception:
            parts.append("")
    text = "\n".join(parts).strip()

    if not text:
        # A scanned circular is the common case here, and naming the cause tells the
        # person what to do instead.
        raise ExtractionFailed(
            "No text could be read from that PDF. It looks like a scan, and text "
            "recognition is not configured — paste the passage you are asking about."
        )
    return Attachment(name, kind, len(payload), text[:MAX_CHARS], pages=len(reader.pages))


def _docx(name: str, kind: str, payload: bytes) -> Attachment:
    try:
        import docx
    except ImportError as exc:
        raise ExtractionFailed("Word reading is unavailable on this server.") from exc

    try:
        document = docx.Document(io.BytesIO(payload))
    except Exception as exc:
        raise ExtractionFailed("That Word file could not be opened.") from exc

    parts = [p.text for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    text = "\n".join(parts).strip()
    if not text:
        raise ExtractionFailed("That Word file had no readable text.")
    return Attachment(name, kind, len(payload), text[:MAX_CHARS])


def as_passages(attachment: Attachment, target_chars: int = 900) -> list[str]:
    """Split extracted text into passages the answer path can cite.

    Split on blank lines first so a clause keeps its own paragraph; only fall back to
    hard slicing when a paragraph is longer than a passage should be. A clause cut in
    half mid-sentence is not quotable, and a quotation is the whole point.
    """
    chunks: list[str] = []
    buffer = ""
    for para in (p.strip() for p in attachment.text.split("\n\n")):
        if not para:
            continue
        if len(buffer) + len(para) + 2 <= target_chars:
            buffer = f"{buffer}\n\n{para}".strip()
            continue
        if buffer:
            chunks.append(buffer)
        while len(para) > target_chars:
            cut = para.rfind(". ", 0, target_chars)
            cut = cut + 1 if cut > target_chars // 2 else target_chars
            chunks.append(para[:cut].strip())
            para = para[cut:].strip()
        buffer = para
    if buffer:
        chunks.append(buffer)
    return [c for c in chunks if c][:24]
